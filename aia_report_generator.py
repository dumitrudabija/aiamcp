"""
AIA Report Generator Module

Handles generation and export of AIA (Algorithmic Impact Assessment) reports.
Extracted from server.py to reduce complexity and improve code organization.

This module mirrors osfi_e23_report_generators.py and provides:
- AIA assessment report generation
- Executive summary creation
- Report formatting and export to Microsoft Word (.docx)
"""

from typing import Dict, Any, List, Optional
from datetime import datetime
from docx import Document
import os
import logging

logger = logging.getLogger(__name__)


class AIAReportGenerator:
    """
    Generates professional AIA compliance reports in Microsoft Word format.
    
    Handles report structure, formatting, and export for Canada's
    Algorithmic Impact Assessment framework.
    """

    def __init__(self, aia_data_extractor):
        """
        Initialize AIA report generator.
        
        Args:
            aia_data_extractor: AIADataExtractor instance for data extraction
        """
        self.aia_data_extractor = aia_data_extractor

    def _get_design_phase_questions(self) -> List[Dict[str, Any]]:
        """Get design phase questions - delegates to data extractor's processor."""
        return self.aia_data_extractor.aia_processor.get_design_phase_questions()

    def _get_assessment_results_for_export(self, session: Dict[str, Any], framework_type: str) -> Optional[Dict[str, Any]]:
        """Extract assessment results from session for export tools."""
        tool_results = session.get("tool_results", {})

        if framework_type == "aia":
            # Look for AIA assessment results
            if "assess_project" in tool_results:
                return tool_results["assess_project"]["result"]
            elif "functional_preview" in tool_results:
                return tool_results["functional_preview"]["result"]
            elif "analyze_project_description" in tool_results:
                return tool_results["analyze_project_description"]["result"]

        elif framework_type == "osfi_e23":
            # Look for OSFI E-23 assessment results
            if "assess_model_risk" in tool_results:
                return tool_results["assess_model_risk"]["result"]
            elif "generate_risk_rating" in tool_results:
                return tool_results["generate_risk_rating"]["result"]
            elif "create_compliance_framework" in tool_results:
                return tool_results["create_compliance_framework"]["result"]

        return None


    def _generate_executive_summary(self, score: int, impact_level: str, project_description: str) -> str:
        """Generate executive summary based on assessment results."""
        description_lower = project_description.lower()

        # Determine risk level
        if score >= 56:
            risk_level = "high risk"
            summary_start = "This system presents significant algorithmic impact risks"
        elif score >= 31:
            risk_level = "moderate risk"
            summary_start = "This system presents moderate algorithmic impact risks"
        else:
            risk_level = "lower risk"
            summary_start = "This system presents relatively low algorithmic impact risks"

        # Add system characteristics
        characteristics = []
        if any(term in description_lower for term in ['ai', 'machine learning', 'neural']):
            characteristics.append("AI/ML-powered")
        if any(term in description_lower for term in ['financial', 'loan', 'credit']):
            characteristics.append("financial decision-making")
        if any(term in description_lower for term in ['automated', 'automatic']):
            characteristics.append("automated processing")
        if any(term in description_lower for term in ['personal', 'sensitive', 'private']):
            characteristics.append("personal data usage")

        char_text = ", ".join(characteristics) if characteristics else "automated decision-making"

        # Generate compliance guidance
        if score >= 56:
            compliance = "Comprehensive governance framework, qualified oversight, and extensive stakeholder consultation are required."
        elif score >= 31:
            compliance = "Enhanced oversight procedures, regular monitoring, and stakeholder engagement processes are recommended."
        else:
            compliance = "Standard operational procedures with basic monitoring and documentation are sufficient."

        return f"{summary_start} due to its {char_text} capabilities. {compliance} The assessment indicates {impact_level} classification under Canada's Algorithmic Impact Assessment framework."


    def _get_assessment_disclaimer(self, assessment_results: Dict[str, Any]) -> str:
        """Get appropriate disclaimer based on assessment type."""
        if 'disclaimer' in assessment_results:
            return assessment_results['disclaimer']
        elif 'functional_risk_score' in assessment_results:
            return "⚠️ Early Indicator - Not Official Assessment. Based on functional characteristics only. Final assessment requires complete stakeholder input and official AIA process completion."
        else:
            return "This assessment is based on available project information and automated analysis. Final AIA compliance requires complete stakeholder input and official government review process."


    def _strip_markdown_formatting(self, text: str) -> str:
        """Strip markdown formatting from text for Word documents."""
        import re
        if not isinstance(text, str):
            return str(text)

        # Remove markdown formatting
        text = re.sub(r'\*\*\*(.*?)\*\*\*', r'\1', text)  # Bold italic
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)      # Bold
        text = re.sub(r'\*(.*?)\*', r'\1', text)          # Italic
        text = re.sub(r'`(.*?)`', r'\1', text)            # Code
        text = re.sub(r'#{1,6}\s+', '', text)             # Headers
        text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)  # Links

        return text.strip()


    def _export_assessment_report(self, arguments: Dict[str, Any]) -> Dict[str, Any]:
        """Export AIA assessment results to a Microsoft Word document."""
        project_name = arguments.get("project_name", "")
        project_description = arguments.get("project_description", "")
        assessment_results = arguments.get("assessment_results", {})
        custom_filename = arguments.get("custom_filename")

        logger.info(f"Exporting assessment report for project: {project_name}")

        # CRITICAL FIX: Validate that assessment_results contains required data
        # Prevent generating misleading reports with default/incomplete values
        if not assessment_results or len(assessment_results) == 0:
            return {
                "error": "export_failed",
                "reason": "Cannot export AIA report: assessment_results is empty or missing",
                "required_action": "Execute 'assess_project' or 'functional_preview' tool first to generate assessment data",
                "workflow_guidance": "If using workflow, the system should auto-inject results. This error indicates no assessment has been completed.",
                "critical_warning": "⚠️ COMPLIANCE RISK: Exporting without assessment data would create misleading documents with incomplete or default values"
            }

        # Check for minimum required assessment fields (score or impact_level)
        has_score = "score" in assessment_results or "functional_risk_score" in assessment_results
        has_impact_level = "impact_level" in assessment_results

        if not has_score and not has_impact_level:
            return {
                "error": "export_failed",
                "reason": "Cannot export AIA report: assessment_results missing required assessment fields",
                "missing_fields": {
                    "score_missing": not has_score,
                    "impact_level_missing": not has_impact_level
                },
                "required_fields": ["score (or functional_risk_score)", "impact_level"],
                "required_action": "Execute 'assess_project' or 'functional_preview' tool to generate complete assessment",
                "received_data": list(assessment_results.keys()) if assessment_results else [],
                "critical_warning": "⚠️ COMPLIANCE RISK: Incomplete assessment data cannot produce valid regulatory documents"
            }
        
        try:
            # Create AIA_Assessments directory if it doesn't exist
            assessments_dir = "./AIA_Assessments"
            os.makedirs(assessments_dir, exist_ok=True)
            
            # Generate filename
            current_date = datetime.now().strftime("%Y-%m-%d")
            if custom_filename:
                filename = f"{custom_filename}.docx"
            else:
                # Clean project name for filename
                clean_project_name = "".join(c for c in project_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
                clean_project_name = clean_project_name.replace(' ', '_')
                filename = f"AIA_Report_{clean_project_name}_{current_date}.docx"
            
            file_path = os.path.join(assessments_dir, filename)
            
            # Create Word document
            doc = Document()
            
            # Add title
            title = doc.add_heading('AIA Assessment Report', 0)
            title.alignment = 1  # Center alignment
            
            # Add project info
            doc.add_heading('Project Information', level=1)
            doc.add_paragraph(f'Project: {project_name}')
            doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
            
            # Extract assessment data
            score = self._extract_score(assessment_results)
            impact_level = self._extract_impact_level(assessment_results)
            design_phase_questions = self._get_design_phase_questions()
            max_score = sum(q['max_score'] for q in design_phase_questions)
            
            doc.add_paragraph(f'Impact Level: {impact_level}')
            doc.add_paragraph(f'Score: {score}/{max_score} points')
            
            # Add executive summary
            doc.add_heading('Executive Summary', level=1)
            summary = self._generate_executive_summary(score, impact_level, project_description)
            doc.add_paragraph(self._strip_markdown_formatting(summary))

            # Add key findings
            doc.add_heading('Key Findings', level=1)
            findings = self._extract_key_findings(assessment_results, project_description)
            for finding in findings:
                p = doc.add_paragraph(self._strip_markdown_formatting(finding))
                p.style = 'List Bullet'

            # Add recommendations
            doc.add_heading('Recommendations', level=1)
            recommendations = self._extract_recommendations(assessment_results, score, impact_level)
            for recommendation in recommendations:
                p = doc.add_paragraph(self._strip_markdown_formatting(recommendation))
                p.style = 'List Bullet'

            # Add project description
            doc.add_heading('Project Description', level=1)
            doc.add_paragraph(self._strip_markdown_formatting(project_description))
            
            # Add disclaimer
            doc.add_heading('Disclaimer', level=1)
            disclaimer_text = self._get_assessment_disclaimer(assessment_results)
            doc.add_paragraph(disclaimer_text)
            
            # Save document
            doc.save(file_path)
            
            # Get file size
            file_size = os.path.getsize(file_path)
            file_size_kb = round(file_size / 1024, 1)
            
            return {
                "assistant_directive": {
                    "critical_instruction": "The COMPLETE AIA compliance report has been generated and saved by the MCP server. Present ONLY the file path and success message below. Do NOT generate, create, or write any additional report content. Do NOT offer to create summaries or additional documents. The Word document is complete and ready for professional review."
                },
                "success": True,
                "file_path": file_path,
                "file_size": f"{file_size_kb}KB",
                "message": f"✅ Canada's AIA compliance report saved successfully to {filename}"
            }
            
        except PermissionError:
            return {
                "success": False,
                "error": "Permission denied - unable to create file. Check write permissions for the directory.",
                "file_path": None
            }
        except OSError as e:
            return {
                "success": False,
                "error": f"File system error: {str(e)}",
                "file_path": None
            }
        except Exception as e:
            logger.error(f"Error creating assessment report: {str(e)}")
            return {
                "success": False,
                "error": f"Failed to create assessment report: {str(e)}",
                "file_path": None
            }
    
    # AIA Data Extraction Methods (delegate to AIADataExtractor)


