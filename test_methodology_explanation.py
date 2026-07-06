#!/usr/bin/env python3
"""
Test the "How to Interpret This Assessment" educational section.

Verifies:
1. The fixed template section is included in generated OSFI E-23 reports by default.
2. It can be disabled via include_methodology_explanation=False.
3. Existing scoring/dimension data flows through unaffected.
"""

import os
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from osfi_e23_report_generators import generate_osfi_e23_report

REQUIRED_PHRASES = [
    "AI does not require a separate risk-management universe",
    "Risk Treatment Outcomes",
    "Residual Risk",
    "Documentation Gaps",
    "Governance Interpretation",
]

MOCK_ASSESSMENT_RESULTS = {
    "risk_score": 65,
    "risk_level": "Medium",
    "dimension_assessments": {
        "misuse_unintended_harm": {"name": "Misuse & Unintended Harm Potential", "risk_level": "Medium"},
    },
    "factor_scores": {},
    "validated_extraction": {},
}


def _generate_and_read(include_methodology_explanation=True):
    doc = Document()
    doc = generate_osfi_e23_report(
        project_name="Methodology Section Test Model",
        project_description="Test project description for methodology section verification.",
        assessment_results=MOCK_ASSESSMENT_RESULTS,
        doc=doc,
        current_stage="design",
        include_methodology_explanation=include_methodology_explanation,
    )
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    doc = Document(tmp_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    os.unlink(tmp_path)
    return full_text


def test_methodology_section_included_by_default():
    full_text = _generate_and_read()

    missing = [phrase for phrase in REQUIRED_PHRASES if phrase not in full_text]
    assert not missing, f"Missing required phrases: {missing}"

    assert "How to Interpret This Assessment" in full_text
    print("PASS: methodology section included by default with all required phrases")


def test_methodology_section_appears_once():
    full_text = _generate_and_read()

    # Should appear exactly once in the main body - not repeated per Annex A dimension table.
    count = full_text.count("How to Interpret This Assessment")
    assert count == 1, f"Expected section heading exactly once, found {count}"
    print("PASS: methodology section appears exactly once (not duplicated in appendix)")


def test_methodology_section_can_be_disabled():
    full_text = _generate_and_read(include_methodology_explanation=False)

    assert "How to Interpret This Assessment" not in full_text
    for phrase in REQUIRED_PHRASES:
        assert phrase not in full_text, f"Unexpected phrase present when disabled: {phrase}"
    print("PASS: methodology section can be disabled via include_methodology_explanation=False")


def test_existing_sections_unaffected():
    full_text = _generate_and_read()

    assert "EXECUTIVE SUMMARY" in full_text
    assert "RISK ASSESSMENT BY DIMENSION" in full_text
    assert "STAGE REQUIREMENTS" in full_text
    assert "ANNEX A" in full_text
    assert "ANNEX B" in full_text
    print("PASS: existing report sections still present and unaffected")


if __name__ == "__main__":
    tests = [
        test_methodology_section_included_by_default,
        test_methodology_section_appears_once,
        test_methodology_section_can_be_disabled,
        test_existing_sections_unaffected,
    ]
    failures = 0
    for test in tests:
        try:
            test()
        except AssertionError as e:
            failures += 1
            print(f"FAIL: {test.__name__}: {e}")

    print(f"\n{len(tests) - failures}/{len(tests)} tests passed")
    sys.exit(0 if failures == 0 else 1)
