"""
OSFI E-23 Report Generation - v4.1 Executive Structure

Generates concise, executive-friendly OSFI E-23 compliance reports:
1. Context and Assessment Summary (incl. 1.2 Model Classification Summary)
2. Risk Rating Methodology and Results (incl. 2.2 Model Type Classification
   Reference - the Level 1-5 taxonomy)
3. Required Actions for Governance Review
Annex A. Official OSFI E-23 Reference (Outcomes, Principles, lifecycle
   definition/components, Appendix 1 model inventory tracking - official
   text only)
Annex B. Fit With Enterprise Risk Management
Annex C. Detailed Question-by-Question Evidence
Annex D. Configurable Governance Matrix (methodology-generated
   implementation guidance, not official OSFI E-23 text)
Annex E. Detailed Model Type Classification Evidence (per-check capability
   evidence detail - not OSFI E-23 lifecycle promotion gates)

This is a presentation/structure layer only - it reorganizes and formats
data that is computed elsewhere (risk_dimension_extraction.py,
model_type_classification.py, conditional_modules.py, osfi_e23_workflow.py,
osfi_e23_structure.py). It does not compute risk scores, does not decide
Capability Evidence Pack triggers, and does not alter the lifecycle
governance matrix values - it only reads and renders them.

Terminology note: model-type classification (Section 1.2, 2.2, Annex E)
uses internal capability evidence checks to determine a model-type level
(1-5). These are methodology-specific classification checks, never called
"promotion gates" in report text - they are unrelated to OSFI E-23
lifecycle-stage promotion/approval (Section 3, Annex A/D).
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import Dict, Any, List, Optional
import logging

from osfi_e23_structure import (
    LIFECYCLE_REQUIREMENTS_BY_RISK,
    get_lifecycle_requirements_for_risk_level,
    get_lifecycle_requirements_comparison,
    OSFI_LIFECYCLE_COMPONENTS,
    OSFI_PRINCIPLES,
    OSFI_OUTCOMES,
    APPENDIX_1_REQUIRED_FIELDS,
    APPENDIX_1_OPTIONAL_FIELDS,
    APPENDIX_1_STAGE_SPECIFIC,
)
from osfi_e23_risk_dimensions import (
    RISK_DIMENSIONS,
    DIMENSION_ORDER,
    get_dimension,
    get_dimension_factors
)

logger = logging.getLogger(__name__)

_RISK_COLORS = {
    "Critical": RGBColor(192, 0, 0),
    "High": RGBColor(255, 102, 0),
    "Medium": RGBColor(255, 192, 0),
    "Low": RGBColor(0, 128, 0),
}

_LIFECYCLE_STAGES = ["design", "review", "deployment", "monitoring", "decommission"]

_PACK_DISPLAY_NAMES = {
    "knowledge_access": "Knowledge Access Pack",
    "action_execution": "Action Execution Pack",
    "autonomy": "Autonomy Pack",
    "vendor_platform": "Vendor / Platform Pack",
}

_GATE_DISPLAY_NAMES = {
    "genai_generation": "GenAI Generation",
    "runtime_retrieval": "Runtime Retrieval for GenAI Grounding",
    "tool_or_action_execution": "Tool / Action Execution",
    "autonomous_operation": "Autonomous Operation",
}

# Source-type labels for report section metadata. Every section/annex that
# reads from model_type_classification.py, conditional_modules.py, or the
# configurable lifecycle matrix must declare which of these it is, so
# readers can tell official OSFI E-23 text apart from methodology
# interpretation, institution-configurable guidance, and model-specific
# evidence/automated output.
_SOURCE_LABEL_TEXT = {
    "official_osfi_e23": "Official OSFI E-23 guideline text",
    "methodology_interpretation": "Methodology interpretation (not official OSFI E-23 text)",
    "institution_configurable": "Institution-configurable implementation guidance (not official OSFI E-23 text)",
    "model_evidence": "Model-specific evidence extracted from the project description",
    "automated_assessment_output": "Automated assessment output (deterministic scoring)",
}


def _add_source_label(doc: Document, *source_types: str):
    """Render a small 'Source: ...' tag under a heading, per the allowed source_type values."""
    label = " + ".join(_SOURCE_LABEL_TEXT.get(s, s) for s in source_types)
    p = doc.add_paragraph()
    run = p.add_run(f"Source: {label}.")
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(89, 89, 89)
    p.paragraph_format.space_after = Pt(6)


# =============================================================================
# TOP-LEVEL ORCHESTRATION
# =============================================================================

def generate_osfi_e23_report(
    project_name: str,
    project_description: str,
    assessment_results: Dict[str, Any],
    doc: Document,
    current_stage: str = "design",
    include_methodology_explanation: bool = True,
    include_model_type_section: bool = True,
    include_conditional_modules_section: bool = True,
    include_governance_matrix: bool = True
) -> Document:
    """
    Generate the OSFI E-23 compliance report (v4.0 executive structure).

    Args:
        include_methodology_explanation: Whether to include Annex B (Fit
            With Enterprise Risk Management) - default True.
        include_model_type_section: Whether to include Annex E (the detailed
            per-check capability evidence table for model-type classification;
            the Section 1.2 summary row always appears - model type is a
            required subsection) - default True.
        include_conditional_modules_section: Whether to include the 12-question
            Key Questions reference table for each triggered Capability
            Evidence Pack (the pack summary in Section 2.5 always appears) -
            default True.
        include_governance_matrix: Whether to include Annex D (the full
            5-stage x 4-risk-level configurable governance matrix) - default True.
    """
    risk_level = assessment_results.get("risk_level", "Medium")
    risk_score = assessment_results.get("risk_score", 50)
    dimension_assessments = assessment_results.get("dimension_assessments", {})
    factor_scores = assessment_results.get("factor_scores", {})
    validated_extraction = assessment_results.get("validated_extraction", {})
    assessment_date = datetime.now().strftime("%B %d, %Y")
    stage_display = current_stage.capitalize()

    final_result = assessment_results.get("final_result") or {}

    # --- Title page ---
    title = doc.add_heading('OSFI E-23 Model Risk Assessment', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading(project_name, level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # --- Front matter: compact metadata block ---
    _add_metadata_block(doc, project_name, assessment_date, current_stage, stage_display, assessment_results, dimension_assessments, final_result, risk_level)

    # --- Professional validation disclaimer (concise, near the front) ---
    _add_professional_validation_disclaimer(doc)

    doc.add_page_break()

    # ================= SECTION 1: CONTEXT AND ASSESSMENT SUMMARY =================
    doc.add_heading('1. CONTEXT AND ASSESSMENT SUMMARY', level=1)
    _add_section_1_1_functional_description(doc, project_description, assessment_results)
    _add_section_1_2_model_classification(doc, assessment_results, include_model_type_section)
    _add_section_1_3_risk_rating_summary(doc, assessment_results, dimension_assessments, current_stage, stage_display, final_result, risk_score)
    _add_section_1_4_required_validation(doc, assessment_results, dimension_assessments, risk_level, final_result)

    doc.add_page_break()

    # ================= SECTION 2: RISK RATING METHODOLOGY AND RESULTS =================
    doc.add_heading('2. RISK RATING METHODOLOGY AND RESULTS', level=1)
    _add_section_2_1_methodology_overview(doc)
    _add_section_2_2_model_type_classification_reference(doc)
    _add_section_2_3_risk_dimensions(doc, dimension_assessments, factor_scores)
    _add_section_2_4_scoring_logic(doc, risk_level, risk_score, dimension_assessments)
    _add_section_2_5_capability_evidence_pack_results(doc, assessment_results, include_conditional_modules_section)
    _add_section_2_6_risk_result_interpretation(doc, assessment_results, dimension_assessments, final_result)

    doc.add_page_break()

    # ================= SECTION 3: REQUIRED ACTIONS FOR GOVERNANCE REVIEW =================
    doc.add_heading('3. REQUIRED ACTIONS FOR GOVERNANCE REVIEW', level=1)
    _add_section_3_disclaimer(doc)
    _add_source_label(doc, "automated_assessment_output", "institution_configurable")
    _add_section_3_1_current_lifecycle_stage(doc, current_stage, stage_display, risk_level)
    _add_section_3_2_stage_governance_requirements(doc, current_stage, stage_display, risk_level)
    _add_section_3_3_required_actions(doc, assessment_results)
    _add_section_3_4_required_validations(doc, assessment_results, dimension_assessments, risk_level, final_result)

    doc.add_page_break()

    # ================= ANNEX A: OFFICIAL OSFI E-23 REFERENCE =================
    _add_annex_a_official_osfi_reference(doc)
    doc.add_page_break()

    # ================= ANNEX B: FIT WITH ENTERPRISE RISK MANAGEMENT =================
    if include_methodology_explanation:
        _add_annex_b_erm_fit(doc)
        doc.add_page_break()

    # ================= ANNEX C: DETAILED QUESTION-BY-QUESTION EVIDENCE =================
    _add_annex_c_detailed_evidence(doc, dimension_assessments, factor_scores, validated_extraction)
    doc.add_page_break()

    # ================= ANNEX D: CONFIGURABLE GOVERNANCE MATRIX =================
    if include_governance_matrix:
        _add_annex_d_configurable_governance_matrix(doc)
        doc.add_page_break()

    # ================= ANNEX E: DETAILED MODEL TYPE CLASSIFICATION EVIDENCE =================
    if include_model_type_section:
        _add_annex_e_model_type_classification_evidence(doc, assessment_results)

    return doc


# =============================================================================
# FRONT MATTER
# =============================================================================

def _add_metadata_block(doc: Document, project_name: str, assessment_date: str, current_stage: str,
                        stage_display: str, assessment_results: Dict[str, Any], dimension_assessments: Dict[str, Any],
                        final_result: Dict[str, Any], risk_level: str):
    """Compact metadata block, as the first thing a reviewer sees. Reports only
    tool-supported/methodology-supported concepts - no final status, blocker,
    condition, or readiness terminology (see module docstring)."""
    classification = assessment_results.get("model_type_classification") or {}
    delivery_model = assessment_results.get("delivery_model") or {}
    triggered_packs = assessment_results.get("capability_evidence_packs", {}).get("triggered", [])
    base_risk_level = final_result.get("base_risk_level", assessment_results.get("risk_level", "N/A"))
    base_risk_score = assessment_results.get("risk_score", 0)

    evidence_gaps_count = len(final_result.get("evidence_gaps", []))
    required_actions_count = len(assessment_results.get("required_governance_actions", []))
    required_validations_count = len(_build_required_validation_rows(assessment_results, dimension_assessments, risk_level, final_result))

    p = doc.add_paragraph()
    p.add_run('ASSESSMENT METADATA').bold = True
    p.paragraph_format.space_after = Pt(6)

    rows = [
        ("Model name", project_name),
        ("Assessment date", assessment_date),
        ("Current lifecycle stage", stage_display),
        ("Base risk score", f"{base_risk_score}/100"),
        ("Base risk rating", base_risk_level),
        ("Model type level", f"Level {classification.get('final_level', 'N/A')}"),
        ("Model type label", classification.get("final_label", "Not classified")),
        ("Delivery model", delivery_model.get("label", "unknown").replace("_", " ").title()),
        ("Triggered capability evidence packs", ", ".join(_PACK_DISPLAY_NAMES.get(p["pack_id"], p["pack_id"]) for p in triggered_packs) or "None"),
        ("Evidence gaps count", str(evidence_gaps_count)),
        ("Required actions count", str(required_actions_count)),
        ("Required validations count", str(required_validations_count)),
    ]
    _add_two_col_or_three_col_table(doc, ['Field', 'Value'], rows)
    doc.add_paragraph()


def _add_professional_validation_disclaimer(doc: Document):
    """Concise professional validation disclaimer, kept near the front."""
    p = doc.add_paragraph()
    run = p.add_run('PROFESSIONAL VALIDATION REQUIRED')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    run = p.add_run(
        "This assessment is an automated preliminary analysis based on the OSFI E-23 model risk "
        "framework. All findings, risk ratings, required actions, and governance recommendations must "
        "be validated by qualified model risk management, compliance, legal, and senior governance "
        "stakeholders before regulatory or production use."
    )
    run.italic = True
    run.font.size = Pt(9)


# =============================================================================
# SECTION 1: CONTEXT AND ASSESSMENT SUMMARY
# =============================================================================

def _add_section_1_1_functional_description(doc: Document, project_description: str, assessment_results: Dict[str, Any]):
    doc.add_heading('1.1 Functional Description', level=2)

    model_type_evidence = assessment_results.get("model_type_classification", {}).get("_evidence") or {}
    # model_type_classification's rationale/gates carry the extracted evidence
    # fields indirectly; the raw evidence dict itself isn't stored on
    # assessment_results, so this description is built from what IS exposed:
    # the classification/delivery_model objects and their gate evidence.
    classification = assessment_results.get("model_type_classification") or {}
    delivery_model = assessment_results.get("delivery_model") or {}
    gates = classification.get("promotion_gates", {})

    parts = []

    tool_gate = gates.get("tool_or_action_execution", {})
    if tool_gate.get("verified"):
        parts.append(
            "Based on available evidence, this system can trigger or cause actions (e.g. system-of-record "
            "changes, tool/API calls, or predefined downstream workflows), not purely informational output."
        )
    elif tool_gate.get("evidence") or tool_gate.get("missing_evidence"):
        parts.append(
            "Based on available evidence, this system's output informs decisions rather than directly "
            "triggering action."
        )
    else:
        parts.append("Whether the system's output only informs decisions or can trigger action is not confirmed.")

    genai_gate = gates.get("genai_generation", {})
    retrieval_gate = gates.get("runtime_retrieval", {})
    if genai_gate.get("verified") and retrieval_gate.get("verified"):
        parts.append("It uses generative AI grounded by runtime retrieval from knowledge sources.")
    elif genai_gate.get("verified"):
        parts.append("It uses generative AI (LLM/foundation model) capability.")
    elif classification.get("final_level") == 1:
        parts.append("Based on available evidence, it uses traditional statistical/ML methods rather than generative AI.")
    else:
        parts.append("Whether it uses generative AI is not confirmed.")

    label = delivery_model.get("label", "unknown")
    if label != "unknown":
        parts.append(f"Delivery model: {label.replace('_', ' ')}.")
    else:
        parts.append("Whether the underlying capability is internally built, vendor-hosted, or embedded in a SaaS product is not confirmed.")

    p = doc.add_paragraph()
    p.add_run(
        "Not confirmed indicates the evidence did not state this explicitly - it is not inferred."
    ).italic = True
    p.paragraph_format.space_after = Pt(6)
    doc.add_paragraph(" ".join(parts))


def _add_section_1_2_model_classification(doc: Document, assessment_results: Dict[str, Any], include_detail: bool):
    doc.add_heading('1.2 Model Classification Summary', level=2)
    _add_source_label(doc, "methodology_interpretation")

    p = doc.add_paragraph()
    p.add_run(
        "Model type is determined by verified capabilities, not product labels. It is an interpretation "
        "lens used to understand the nature of the model and the relevant evidence packs. It is not a "
        "separate risk score and does not replace lifecycle-stage governance approval."
    ).italic = True
    p.paragraph_format.space_after = Pt(8)

    classification = assessment_results.get("model_type_classification") or {}
    delivery_model = assessment_results.get("delivery_model") or {}
    gates = classification.get("promotion_gates", {})
    verified_gates = [_GATE_DISPLAY_NAMES.get(g, g) for g, v in gates.items() if v.get("verified")]
    not_verified_gates = [_GATE_DISPLAY_NAMES.get(g, g) for g, v in gates.items() if not v.get("verified")]

    rows = [
        ("Final model type", classification.get("final_label", "Not classified")),
        ("Model type level", f"Level {classification.get('final_level', 'N/A')}"),
        ("Delivery model", delivery_model.get("label", "unknown").replace("_", " ").title()),
        ("Classification method", "Deterministic capability evidence checks"),
        ("Classification confidence", str(classification.get("confidence", "N/A")).title()),
        ("Key capabilities verified", ", ".join(verified_gates) or "None"),
        ("Key capabilities not verified", ", ".join(not_verified_gates) or "None"),
        ("Rationale", classification.get("rationale", "N/A")),
    ]
    _add_two_col_or_three_col_table(doc, ['Attribute', 'Value'], rows)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(
        "Model-type classification evidence checks are methodology-specific checks used to identify the "
        "model's capabilities. They are not lifecycle-stage approval gates and do not determine approval, "
        "readiness, or production use."
    ).italic = True

    if include_detail:
        p = doc.add_paragraph()
        p.add_run(
            "Full capability evidence detail (status, evidence, missing evidence, and classification "
            "implication for each check) is provided in Annex E, Detailed Model Type Classification Evidence."
        ).italic = True


def _add_section_1_3_risk_rating_summary(doc: Document, assessment_results: Dict[str, Any], dimension_assessments: Dict[str, Any],
                                         current_stage: str, stage_display: str, final_result: Dict[str, Any],
                                         risk_score: int):
    doc.add_heading('1.3 Risk Rating Summary', level=2)

    base_risk_level = final_result.get("base_risk_level", assessment_results.get("risk_level", "N/A"))
    evidence_gaps = final_result.get("evidence_gaps", [])
    required_actions = assessment_results.get("required_governance_actions") or []
    triggered_packs = assessment_results.get("capability_evidence_packs", {}).get("triggered", [])

    high_risk_dims = [d for d, a in dimension_assessments.items() if a.get("risk_level") in ("High", "Critical")]
    dim_names = [get_dimension(d).get("name", d) for d in high_risk_dims[:3]]

    summary = (
        f"The base risk rating is {base_risk_level} based on the core assessment ({risk_score}/100, "
        f"47-question assessment across 8 risk dimensions). The assessment identified {len(evidence_gaps)} "
        f"evidence gap(s) and {len(required_actions)} required action(s) that should be reviewed by the "
        f"institution's model governance process."
    )
    doc.add_paragraph(summary)

    p = doc.add_paragraph()
    p.add_run(
        "The report does not determine lifecycle-stage approval. Identified required actions should be "
        "reviewed and dispositioned through the institution's lifecycle governance process before any "
        "stage transition decision."
    ).italic = True

    if dim_names:
        doc.add_paragraph(f"Top risk drivers: {', '.join(dim_names)}.")

    doc.add_paragraph(
        f"Triggered Capability Evidence Packs: "
        + (", ".join(_PACK_DISPLAY_NAMES.get(p["pack_id"], p["pack_id"]) for p in triggered_packs) or "None") + "."
    )
    doc.add_paragraph(f"Evidence gaps: {len(evidence_gaps)}. Required actions: {len(required_actions)}.")


def _add_section_1_4_required_validation(doc: Document, assessment_results: Dict[str, Any], dimension_assessments: Dict[str, Any],
                                         risk_level: str, final_result: Dict[str, Any]):
    doc.add_heading('1.4 Required Validation', level=2)

    rows = _build_required_validation_rows(assessment_results, dimension_assessments, risk_level, final_result)
    _add_two_col_or_three_col_table(doc, ['Validation area', 'Required validation', 'Trigger/source', 'Due stage'], rows)


def _build_required_validation_rows(assessment_results: Dict[str, Any], dimension_assessments: Dict[str, Any],
                                    risk_level: str, final_result: Dict[str, Any]) -> List[tuple]:
    triggered_ids = {p["pack_id"] for p in assessment_results.get("capability_evidence_packs", {}).get("triggered", [])}
    evidence_gaps = final_result.get("evidence_gaps", [])

    def dim_risk(dim_id: str) -> str:
        return dimension_assessments.get(dim_id, {}).get("risk_level", "Not Assessed")

    rows = [
        ("Independent model risk review", "Required", "Base risk assessment", "Deployment"),
        ("Lifecycle-stage approval", "Required", "Lifecycle governance matrix", "Stage transition"),
    ]
    if evidence_gaps:
        rows.append(("Evidence gap remediation", "Required", f"{len(evidence_gaps)} evidence gap(s) identified", "Deployment"))
    if triggered_ids:
        rows.append((
            "Capability Evidence Pack control validation", "Required",
            ", ".join(_PACK_DISPLAY_NAMES.get(pid, pid) for pid in triggered_ids), "Deployment",
        ))
    if "vendor_platform" in triggered_ids:
        rows.append(("Vendor risk review", "Required", "Vendor / Platform Pack triggered", "Deployment"))
    if "action_execution" in triggered_ids:
        rows.append(("Action logging validation", "Required", "Action Execution Pack triggered", "Deployment"))
        rows.append(("Rollback or compensating-control validation", "Required", "Action Execution Pack triggered", "Deployment"))
    if "autonomy" in triggered_ids:
        rows.append(("Monitoring and incident response validation", "Required", "Autonomy Pack triggered", "Deployment"))
    if dim_risk("fairness_customer_impact") in ("Medium", "High", "Critical"):
        rows.append(("Fairness/bias review", "Required", f"Fairness & Customer Impact: {dim_risk('fairness_customer_impact')}", "Deployment"))
        rows.append(("Legal/compliance review", "Required", f"Fairness & Customer Impact: {dim_risk('fairness_customer_impact')}", "Deployment"))
    if dim_risk("data_provenance_supply_chain") in ("Medium", "High", "Critical"):
        rows.append(("Privacy review", "Required", f"Data Provenance & Supply Chain Risk: {dim_risk('data_provenance_supply_chain')}", "Deployment"))
    if dim_risk("operational_security") in ("Medium", "High", "Critical"):
        rows.append(("Security review", "Required", f"Operational & Security Risk: {dim_risk('operational_security')}", "Deployment"))
    if risk_level in ("High", "Critical"):
        rows.append(("Senior risk committee or executive approval", "Required", f"{risk_level} risk level", "Deployment"))
    return rows


# =============================================================================
# SECTION 2: RISK RATING METHODOLOGY AND RESULTS
# =============================================================================

def _add_section_2_1_methodology_overview(doc: Document):
    doc.add_heading('2.1 Methodology Overview', level=2)
    _add_source_label(doc, "methodology_interpretation")
    doc.add_paragraph(
        "The base risk score is calculated from the common 47-question assessment across 8 risk "
        "dimensions. By default, dimensions and questions are equally weighted unless institution-specific "
        "configuration overrides are provided. Weights, thresholds, and governance mappings are "
        "configurable to align with institutional risk appetite."
    )
    notes = [
        "The assessment uses 8 risk dimensions and 47 questions.",
        "Each question is assessed using deterministic scoring criteria.",
        "Extracted evidence supports factor values - Claude extracts evidence; deterministic logic scores it.",
        "Missing evidence is treated as an evidence gap (defaulted to Medium risk) rather than assumed low risk.",
        "The default methodology uses equal weighting unless configured otherwise.",
        "Weights, thresholds, and governance rules are tunable by institution.",
    ]
    for n in notes:
        p = doc.add_paragraph(f'• {n}')
        p.paragraph_format.left_indent = Inches(0.25)


def _add_section_2_2_model_type_classification_reference(doc: Document):
    doc.add_heading('2.2 Model Type Classification Reference', level=2)
    _add_source_label(doc, "methodology_interpretation")

    p = doc.add_paragraph()
    p.add_run(
        "The model-type levels below are methodology-specific categories used to interpret model "
        "capabilities and trigger relevant evidence packs. They are not official OSFI E-23 model "
        "categories and do not replace the model risk rating."
    ).italic = True
    p.paragraph_format.space_after = Pt(8)

    rows = [
        (
            "1", "Traditional model / rules-based or statistical model",
            "A model that uses deterministic rules, statistical techniques, scorecards, regression, "
            "decision trees, ensembles, or other non-generative analytical methods to process inputs and "
            "produce scores, classifications, forecasts, decisions, or recommendations.",
            "Traditional ML or statistical model; scorecard, regression, tree-based model, rules engine, "
            "optimization model, or forecasting model; no evidence of LLM or generative AI use; no "
            "runtime GenAI retrieval; no AI-mediated tool or action execution; no autonomous operation.",
            "Assess through the 47-question core. Do not trigger GenAI-specific evidence packs unless "
            "other capabilities are verified.",
        ),
        (
            "2", "Generative AI / LLM output generation",
            "A model or AI system that uses a large language model, foundation model, or generative AI "
            "component to generate, summarize, classify, transform, draft, or explain content, but does "
            "not use runtime retrieval to ground outputs and does not trigger downstream action "
            "execution.",
            "LLM or foundation model use; prompt-based generation; text, code, image, reasoning, "
            "summary, explanation, or classification generation; no verified runtime retrieval for "
            "grounding; no verified system-of-record write or downstream action execution; no verified "
            "autonomous operation.",
            "Assess through the 47-question core. Apply GenAI-relevant factors such as output quality, "
            "hallucination/confabulation risk, explainability, prompt-injection exposure if applicable, "
            "and data leakage controls. Do not trigger Knowledge Access Pack unless runtime retrieval is "
            "verified.",
        ),
        (
            "3", "Knowledge-grounded GenAI / retrieval-augmented AI",
            "A generative AI system that retrieves enterprise, external, or indexed knowledge at "
            "runtime to ground, contextualize, cite, or generate outputs, but does not execute actions "
            "or change system state.",
            "Runtime retrieval used for GenAI grounding; RAG, vector search, enterprise search, "
            "document retrieval, policy retrieval, CRM/SharePoint/Confluence/email/PDF/database "
            "retrieval; retrieved content used to produce or support generated outputs; no verified "
            "action execution; no verified autonomous operation.",
            "Trigger the Knowledge Access Pack. Assess retrieval quality, source ownership, entitlement "
            "enforcement, source freshness, traceability, prompt-injection risk through retrieved "
            "content, and sensitive-data leakage risk.",
        ),
        (
            "4", "Agentic workflow / AI-mediated action execution",
            "An AI system whose outputs can trigger or cause downstream workflows, tool calls, API "
            "calls, system-of-record writes, routing decisions, external communications, approvals, "
            "denials, transactions, ticket creation, or other changes to system state, but where "
            "autonomous AI discretion over continuation, sequencing, or goal pursuit is not verified.",
            "Model output changes system state; model output triggers a predefined workflow; system or "
            "service account has system-of-record write permission; model output initiates external "
            "communication; model output triggers transaction, approval, denial, routing, or case "
            "creation; AI discretion over whether to act, continue, plan, or revise steps is not "
            "verified.",
            "Trigger the Action Execution Pack. Do not trigger the Autonomy Pack unless autonomous "
            "operation is separately verified. Assess action inventory, action authorization, trigger "
            "criteria, audit logging, rollback or compensating controls, human approval gates, "
            "duplicate-action controls, downstream validation, monitoring, and emergency disablement.",
        ),
        (
            "5", "Autonomous agent",
            "An AI system with verified discretion to decide whether to act or continue, select tools "
            "or actions, sequence steps, revise plans, pursue goals, loop or retry based on outcomes, "
            "use memory or state to continue activity, delegate to other agents or subprocesses, or "
            "otherwise operate without human approval at each step.",
            "AI decides whether to act or continue; AI selects next step based on intermediate results; "
            "dynamic multi-step planning; goal pursuit across actions; adaptive plan revision; looping "
            "or retry based on outcomes; memory or state-driven continuation; delegation to other "
            "agents or subprocesses; no human approval required at each step.",
            "Trigger the Autonomy Pack. If action execution is also verified, trigger the Action "
            "Execution Pack as well. Assess stop conditions, kill switch, scope boundaries, plan "
            "traceability, loop/retry limits, human oversight model, guardrail testing, memory/state "
            "controls, autonomous monitoring, delegation controls, and autonomy-specific governance "
            "approval.",
        ),
    ]
    _add_two_col_or_three_col_table(
        doc, ['Level', 'Label', 'Description', 'Typical evidence', 'Methodology implication'], rows
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Level Assignment Rules').bold = True
    p.paragraph_format.space_after = Pt(4)

    rules = [
        "Assign the highest level whose required capability evidence is verified.",
        "Do not assign a level based on product labels, vendor marketing, or generic AI language.",
        "If evidence is missing, mark the capability as \"not verified\" rather than inferring it.",
        "Level 4 requires verified action execution or system-state change.",
        "Level 5 requires verified autonomous discretion over action, continuation, sequencing, "
        "planning, or goal pursuit.",
        "Scheduled workflows, batch jobs, event triggers, straight-through processing, or no human "
        "review per transaction do not by themselves establish Level 5 autonomy.",
        "A system can be Level 4 without being Level 5.",
        "Capability Evidence Packs are triggered by verified capabilities, not by the level label alone.",
        "Model type does not change the base risk score directly unless institution-specific "
        "configuration says otherwise.",
        "Model type is used to interpret risk, trigger evidence packs, and guide governance questions.",
    ]
    for i, rule in enumerate(rules, 1):
        p = doc.add_paragraph(f"{i}. {rule}")
        p.paragraph_format.left_indent = Inches(0.25)


def _add_section_2_3_risk_dimensions(doc: Document, dimension_assessments: Dict[str, Any], factor_scores: Dict[str, Any]):
    doc.add_heading('2.3 Risk Dimensions', level=2)

    rows = []
    for dim_id in DIMENSION_ORDER:
        dim_info = get_dimension(dim_id)
        if not dim_info:
            continue
        dim_assessment = dimension_assessments.get(dim_id, {})
        evidence_summary = _first_available_evidence(factor_scores.get(dim_id, []))
        rows.append((
            dim_info.get("name", dim_id),
            dim_info.get("core_question", ""),
            dim_assessment.get("risk_level", "Not Assessed"),
            evidence_summary,
        ))
    _add_two_col_or_three_col_table(doc, ['Risk dimension', 'Core question', 'Assessment result', 'Key evidence summary'], rows)


def _first_available_evidence(dim_factor_scores: List[Dict[str, Any]]) -> str:
    for fs in dim_factor_scores:
        evidence = fs.get("evidence")
        if evidence:
            return evidence
    return "No evidence summary available."


def _add_section_2_4_scoring_logic(doc: Document, risk_level: str, risk_score: int, dimension_assessments: Dict[str, Any]):
    doc.add_heading('2.4 Scoring Logic', level=2)

    p = doc.add_paragraph()
    p.add_run(f'Base score: ').bold = True
    p.add_run(f'{risk_score}/100 -> {risk_level.upper()}')
    p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "Risk level thresholds (0-100 scale, derived from the underlying 1-4 factor scoring scale): "
        "below 37.5 = Low, 37.5-62.4 = Medium, 62.5-87.4 = High, 87.5 and above = Critical."
    )
    doc.add_paragraph(
        "Each of the 8 dimensions is scored as the average of its factors' numeric scores (see 2.2 Risk "
        "Dimensions for dimension-level results). Factors not stated in the project description default to "
        "Medium risk and are tracked as evidence gaps rather than assumed to be low risk. Factors marked "
        "Not Applicable are scored per their configured not-applicable risk level. A small number of "
        "portfolio-level factors are excluded from the dimension average entirely when institution-wide "
        "inventory data is unavailable, and tracked separately as follow-up actions."
    )
    p = doc.add_paragraph()
    p.add_run(
        "Capability Evidence Packs do not create separate risk scores. Pack findings can create evidence "
        "gaps and required actions, but they do not directly change the 47-question base risk score "
        "unless explicitly configured."
    ).bold = True

    not_stated_total = sum(a.get("not_stated_count", 0) for a in dimension_assessments.values())
    if not_stated_total:
        doc.add_paragraph(f"{not_stated_total} factor(s) across all dimensions were not stated in the project description and defaulted to Medium risk.")


def _add_section_2_5_capability_evidence_pack_results(doc: Document, assessment_results: Dict[str, Any], include_key_questions: bool):
    doc.add_heading('2.5 Capability Evidence Pack Results', level=2)
    _add_source_label(doc, "methodology_interpretation", "model_evidence")

    p = doc.add_paragraph()
    p.add_run(
        "Capability Evidence Packs identify additional evidence needs; they do not create separate risk "
        "scores. Evidence Pack findings are mapped back to the existing risk dimensions. Packs produce "
        "findings, evidence gaps, required actions, and required validation only, and do not produce any "
        "other determination, decision, or status value."
    ).italic = True
    p.paragraph_format.space_after = Pt(8)

    packs_data = assessment_results.get("capability_evidence_packs") or {}
    triggered_packs = {p["pack_id"]: p for p in packs_data.get("triggered", [])}
    not_triggered_packs = {p["pack_id"]: p for p in packs_data.get("not_triggered", [])}

    rows = []
    for pack_id in ["knowledge_access", "action_execution", "autonomy", "vendor_platform"]:
        pack_name = _PACK_DISPLAY_NAMES[pack_id]
        if pack_id in triggered_packs:
            pack = triggered_packs[pack_id]
            mapped_names = pack.get("mapped_dimension_names") or pack.get("mapped_dimensions", [])
            key_questions = pack.get("key_questions") or []
            required_validation = (
                f"{len(key_questions)} item(s) - see Key Questions below" if key_questions else "None"
            )
            rows.append((
                pack_name, "Yes",
                pack.get("trigger_reason", "N/A"),
                ", ".join(mapped_names),
                (pack.get("findings") or ["N/A"])[0],
                "; ".join(pack.get("evidence_gaps", [])) or "None",
                "; ".join(pack.get("required_actions", [])) or "None",
                required_validation,
            ))
        else:
            not_trig = not_triggered_packs.get(pack_id, {})
            rows.append((pack_name, "No", not_trig.get("reason", "Not triggered."), "N/A", "N/A", "N/A", "N/A", "N/A"))

    _add_two_col_or_three_col_table(
        doc,
        ['Pack', 'Triggered', 'Trigger reason', 'Mapped dimensions', 'Finding', 'Evidence gaps', 'Required actions', 'Required validation'],
        rows,
    )

    if include_key_questions and triggered_packs:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(
            "Key Questions for Institutional Follow-Up (reference only - not verified against the project "
            "description; evidence for these external/operational controls should be confirmed internally "
            "with the accountable control owner):"
        ).italic = True
        p.paragraph_format.space_after = Pt(6)

        for pack_id, pack in triggered_packs.items():
            key_questions = pack.get("key_questions") or []
            if not key_questions:
                continue
            doc.add_heading(_PACK_DISPLAY_NAMES.get(pack_id, pack_id), level=3)
            kq_rows = [
                (q["question_id"], q["question"], q["expected_evidence_or_control"], q["condition_if_missing"])
                for q in key_questions
            ]
            _add_two_col_or_three_col_table(
                doc,
                ['ID', 'Question', 'Expected Evidence / Control', 'Required Action'],
                kq_rows,
            )
            doc.add_paragraph()


def _add_section_2_6_risk_result_interpretation(doc: Document, assessment_results: Dict[str, Any], dimension_assessments: Dict[str, Any],
                                                final_result: Dict[str, Any]):
    doc.add_heading('2.6 Risk Result Interpretation', level=2)

    high_risk_dims = [d for d, a in dimension_assessments.items() if a.get("risk_level") in ("High", "Critical")]
    dim_names = [get_dimension(d).get("name", d) for d in high_risk_dims]
    evidence_gaps = final_result.get("evidence_gaps", [])
    required_actions = assessment_results.get("required_governance_actions") or []

    parts = []
    if dim_names:
        parts.append(f"The score is primarily driven by: {', '.join(dim_names)}.")
    else:
        parts.append("No dimension scored High or Critical - the score is driven by a broadly Low-to-Medium risk profile across dimensions.")

    if evidence_gaps:
        parts.append(f"The most material evidence gaps ({len(evidence_gaps)} total) come from triggered Capability Evidence Packs and should be reviewed by the institution's model governance process.")
    else:
        parts.append("No material evidence gaps were identified from triggered Capability Evidence Packs.")

    if required_actions:
        parts.append(f"{len(required_actions)} required action(s) should be reviewed by the institution's model governance process: see 3. Required Actions for Governance Review for the full list.")

    parts.append(
        "Residual concern reflects the evidence gaps and required actions above, combined with the "
        "professional validations in 1.4 that have not yet been completed. Full question-by-"
        "question evidence is provided in Annex C, not duplicated here."
    )
    doc.add_paragraph(" ".join(parts))


# =============================================================================
# SECTION 3: REQUIRED ACTIONS FOR GOVERNANCE REVIEW
# =============================================================================

def _add_section_3_disclaimer(doc: Document):
    p = doc.add_paragraph()
    p.add_run(
        "The items in this section are automated assessment outputs intended to support governance "
        "review. They are not approval decisions and do not by themselves permit or prevent moving to "
        "the next lifecycle stage. The institution's accountable model governance process must review, "
        "accept, modify, or close these actions."
    ).bold = True
    p.paragraph_format.space_after = Pt(8)


def _add_section_3_1_current_lifecycle_stage(doc: Document, current_stage: str, stage_display: str, risk_level: str):
    doc.add_heading('3.1 Current Lifecycle Stage and Base Risk Rating', level=2)

    rows = [
        ("Current lifecycle stage", stage_display),
        ("Base risk rating", risk_level),
    ]
    _add_two_col_or_three_col_table(doc, ['Attribute', 'Value'], rows)


def _add_section_3_2_stage_governance_requirements(doc: Document, current_stage: str, stage_display: str, risk_level: str):
    doc.add_heading('3.2 Stage-Specific Required Actions', level=2)

    stage_requirements = get_lifecycle_requirements_for_risk_level(current_stage, risk_level)
    if not stage_requirements:
        doc.add_paragraph("Requirement not configured.")
        return

    p = doc.add_paragraph()
    p.add_run('GOVERNANCE INTENSITY: ').bold = True
    p.add_run(f'Requirements below are scaled to {risk_level} risk level per OSFI Principle 2.3, for the {stage_display} stage only. Full matrix across all stages and risk levels is in Annex A.')
    p.paragraph_format.space_after = Pt(10)

    checklist_items = _get_checklist_items_for_stage(current_stage)
    rows = []
    for req_name, req_value in stage_requirements.items():
        req_display = req_name.replace('_', ' ').title()
        items = checklist_items.get(req_name, [])
        evidence_expected = "; ".join(items) if items else "Not configured."
        rows.append((req_display, req_value or "Not configured.", evidence_expected, "Not started", "TBD", f"{stage_display} stage completion"))

    _add_two_col_or_three_col_table(doc, ['Requirement area', 'Required for current risk level', 'Evidence expected', 'Action status', 'Owner', 'Due stage'], rows)


_PRIORITY_ORDER = {"blocker": 0, "high": 1, "medium": 2, "low": 3}
# Internal priority values from osfi_e23_workflow.py/conditional_modules.py
# include "blocker" as the most urgent tier - displayed as "Critical" so the
# report never shows the word "Blocker" as a first-class label; the
# underlying priority value/ordering is unchanged.
_PRIORITY_DISPLAY = {"blocker": "Critical", "high": "High", "medium": "Medium", "low": "Low"}

_BASELINE_SOURCE_BY_CATEGORY = {
    "Issue remediation": "Base risk",
    "Workflow / ticket creation": "Evidence gap",
}

# A couple of fixed baseline-action sentences in osfi_e23_workflow.py (out of
# scope for this report-only pass) use "blocker"/"condition" vocabulary
# directly. Rewritten here at render time only - the underlying action list
# in assessment_results is not modified.
_ACTION_TEXT_REWRITES = {
    "Resolve all identified production blockers before deployment.":
        "Resolve all identified required actions before deployment.",
    "Document the classification, 47-question assessment results, risk level, and conditions for audit purposes.":
        "Document the classification, 47-question assessment results, risk level, and required actions for audit purposes.",
}


def _sanitize_action_text(text: str) -> str:
    return _ACTION_TEXT_REWRITES.get(text, text)


def _add_section_3_3_required_actions(doc: Document, assessment_results: Dict[str, Any]):
    doc.add_heading('3.3 Required Actions', level=2)

    p = doc.add_paragraph()
    p.add_run(
        "Required actions in this table come from three sources: the base risk rating and model "
        "inventory requirements (Source = Base risk), triggered Capability Evidence Packs (Source = pack "
        "name), and identified evidence gaps (Source = Evidence gap). Every evidence gap generates at "
        "least one required action. Owner and action status are institution-assignable placeholders."
    ).italic = True
    p.paragraph_format.space_after = Pt(10)

    rows = _build_required_action_rows(assessment_results)
    if not rows:
        doc.add_paragraph("No required actions were generated.")
        return

    _add_two_col_or_three_col_table(doc, ['Priority', 'Category', 'Action', 'Source', 'Owner', 'Due stage', 'Action status'], rows)


def _build_required_action_rows(assessment_results: Dict[str, Any]) -> List[tuple]:
    """
    Build the 3.3 Required Actions display rows. Reads (does not mutate) the
    existing required_governance_actions produced by osfi_e23_workflow.py,
    relabels `source_pack` into a Source column, maps the internal priority
    value to a display label (see _PRIORITY_DISPLAY), and sanitizes the small
    set of known baseline-action strings that use blocker/condition wording
    (see _ACTION_TEXT_REWRITES).
    """
    actions = assessment_results.get("required_governance_actions") or []

    ranked_rows = []
    for action in actions:
        source_pack = action.get("source_pack")
        if source_pack:
            source = _PACK_DISPLAY_NAMES.get(source_pack, source_pack)
        else:
            source = _BASELINE_SOURCE_BY_CATEGORY.get(action.get("category", ""), "Base risk")
        priority_key = action.get("priority", "medium").lower()
        ranked_rows.append((
            _PRIORITY_ORDER.get(priority_key, 2),
            _PRIORITY_DISPLAY.get(priority_key, "Medium"),
            action.get("category", "General"),
            _sanitize_action_text(action.get("action", "")),
            source,
            action.get("owner", "TBD"),
            action.get("due_stage", "TBD"),
            "Open",
        ))

    ranked_rows.sort(key=lambda r: r[0])
    return [r[1:] for r in ranked_rows]


def _add_section_3_4_required_validations(doc: Document, assessment_results: Dict[str, Any], dimension_assessments: Dict[str, Any],
                                          risk_level: str, final_result: Dict[str, Any]):
    doc.add_heading('3.4 Required Validations', level=2)

    rows = _build_required_validation_rows(assessment_results, dimension_assessments, risk_level, final_result)
    if not rows:
        doc.add_paragraph("No required validation items were generated.")
        return

    validation_areas = [r[0] for r in rows]
    doc.add_paragraph(
        f"{len(rows)} required validation item(s) apply to this assessment: {', '.join(validation_areas)}. "
        "See 1.4 Required Validation for full detail (required validation, trigger/source, and due stage)."
    )


# =============================================================================
# SHARED TABLE HELPER
# =============================================================================

# Column-width tuning for _add_two_col_or_three_col_table's content-based sizing.
_COL_WIDTH_FLOOR = Inches(0.55)   # no column shrinks below this, however short its content
_COL_WIDTH_LINE_CAP = 45          # longest-line chars considered per column, so one huge
                                  # paragraph (e.g. Rationale) can't starve every other column


def _add_two_col_or_three_col_table(doc: Document, headers: List[str], rows: List[tuple]):
    """
    Add a bold-header table - supports an arbitrary number of columns.

    Column widths are sized to content (not left at Word's equal-width default):
    each column's width is proportional to the longest single line found in that
    column (header or data, splitting multi-line cells on "\n"), so short fields
    like an ID or risk level get a narrow column and long free-text fields like
    a rationale or evidence summary get the room they need to stay readable.
    """
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.autofit = False
    table.allow_autofit = False

    header_cells = table.rows[0].cells
    for col_idx, header_text in enumerate(headers):
        header_cells[col_idx].text = header_text
        for paragraph in header_cells[col_idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row_idx, row_values in enumerate(rows, 1):
        row_cells = table.rows[row_idx].cells
        for col_idx, value in enumerate(row_values):
            row_cells[col_idx].text = str(value)

    _size_columns_to_content(doc, table, headers, rows)


def _size_columns_to_content(doc: Document, table, headers: List[str], rows: List[tuple]):
    """Set explicit, content-proportional column widths (see module docstring
    on _add_two_col_or_three_col_table for the rationale)."""
    usable_width = _usable_page_width(doc)
    num_cols = len(headers)

    def longest_line(value: Any) -> int:
        return max((len(line) for line in str(value).split("\n")), default=0)

    weights = []
    for col_idx in range(num_cols):
        col_values = [headers[col_idx]] + [row[col_idx] for row in rows if col_idx < len(row)]
        weight = min(max((longest_line(v) for v in col_values), default=1), _COL_WIDTH_LINE_CAP)
        weights.append(max(weight, 1))

    floor_total = _COL_WIDTH_FLOOR * num_cols
    remaining = max(usable_width - floor_total, 0)
    weight_total = sum(weights)

    for col_idx, weight in enumerate(weights):
        extra = int(remaining * (weight / weight_total)) if weight_total else 0
        width = _COL_WIDTH_FLOOR + extra
        # docx.add_table() stamps an explicit width on every individual cell
        # (tcW), which takes priority over the table's tblGrid column width in
        # Word's rendering - the grid alone is not enough, every cell in the
        # column must be set too.
        table.columns[col_idx].width = width
        for row in table.rows:
            row.cells[col_idx].width = width


def _usable_page_width(doc: Document) -> int:
    """Page width minus left/right margins, in EMU (python-docx's internal unit)."""
    section = doc.sections[0]
    return section.page_width - section.left_margin - section.right_margin


def _get_checklist_items_for_stage(stage: str) -> Dict[str, List[str]]:
    """Get checklist items for each requirement area in a stage."""
    checklist_by_stage = {
        "design": {
            "documentation_depth": ["Document model rationale and business purpose", "Record design alternatives considered"],
            "data_quality_assessment": ["Verify data sources meet quality standards", "Document data lineage and provenance"],
            "bias_fairness_analysis": ["Screen for potential bias in training data", "Document fairness considerations"],
            "approval_authority": ["Obtain required design approval", "Document approval in model inventory"],
        },
        "review": {
            "validation_independence": ["Assign independent reviewer", "Document reviewer qualifications"],
            "testing_scope": ["Execute required test suite", "Document test results and findings"],
            "challenger_model": ["Develop challenger model (if required)", "Compare performance against primary model"],
            "explainability_review": ["Validate model explainability", "Document explanation methodology"],
            "approval_authority": ["Obtain validation sign-off", "Document approval for deployment"],
        },
        "deployment": {
            "environment_verification": ["Verify production environment configuration", "Confirm integration points tested"],
            "parallel_run_period": ["Execute parallel run (if required)", "Document parallel run results"],
            "rollback_capability": ["Test rollback procedures", "Document rollback instructions"],
            "human_override_controls": ["Implement override mechanisms", "Document override procedures"],
            "go_live_approval": ["Obtain go-live approval", "Document deployment date and approver"],
        },
        "monitoring": {
            "performance_review_frequency": ["Establish monitoring schedule", "Configure performance dashboards"],
            "drift_monitoring": ["Implement drift detection", "Set drift alert thresholds"],
            "fairness_monitoring": ["Monitor for disparate impact", "Track fairness metrics"],
            "incident_escalation_time": ["Define escalation procedures", "Document escalation contacts"],
            "revalidation_trigger": ["Define revalidation triggers", "Schedule periodic revalidation"],
        },
        "decommission": {
            "retention_period": ["Archive model artifacts per retention policy", "Document retention start date"],
            "documentation_to_retain": ["Compile final documentation package", "Store in approved archive location"],
            "stakeholder_notification": ["Notify all stakeholders of retirement", "Document notification confirmations"],
            "downstream_impact_review": ["Assess downstream system impacts", "Verify no residual dependencies"],
        },
    }
    return checklist_by_stage.get(stage, {})


# =============================================================================
# ANNEX A: OFFICIAL OSFI E-23 REFERENCE
# =============================================================================
# Official OSFI Guideline E-23 reference content ONLY - Outcomes, Principles,
# lifecycle definition/components, and Appendix 1 model inventory tracking
# expectations. Must never include methodology-specific model-type levels,
# capability classification evidence, the configurable governance matrix
# (see Annex D), internal approval authorities, or configured monitoring/
# retention periods.

def _add_annex_a_official_osfi_reference(doc: Document):
    doc.add_heading('ANNEX A. OFFICIAL OSFI E-23 REFERENCE', level=1)
    _add_source_label(doc, "official_osfi_e23")

    p = doc.add_paragraph()
    p.add_run(
        "This annex reproduces official OSFI Guideline E-23 reference content only: the Outcomes, "
        "Principles, model lifecycle definition and components, and Appendix 1 model inventory "
        "information tracking expectations. It does not include methodology-specific model-type "
        "levels, capability classification evidence, or institution-configurable governance detail - "
        "see Annex D and Annex E for that content."
    ).italic = True
    p.paragraph_format.space_after = Pt(10)

    # --- A.1 Outcomes ---
    doc.add_heading('A.1 Outcomes', level=2)
    outcome_rows = list(OSFI_OUTCOMES.items())
    _add_two_col_or_three_col_table(doc, ['Outcome', 'Description'], outcome_rows)
    doc.add_paragraph()

    # --- A.2 Principles ---
    doc.add_heading('A.2 Principles', level=2)
    for outcome_num, outcome_text in OSFI_OUTCOMES.items():
        doc.add_heading(f'Outcome {outcome_num}: {outcome_text}', level=3)
        principle_rows = [
            (principle_num, principle_text)
            for principle_num, principle_text in OSFI_PRINCIPLES.items()
            if principle_num.startswith(f"{outcome_num}.")
        ]
        if principle_rows:
            _add_two_col_or_three_col_table(doc, ['Principle', 'Text'], principle_rows)
        doc.add_paragraph()

    # --- A.3 Model Lifecycle Definition and Components ---
    doc.add_heading('A.3 Model Lifecycle Definition and Components', level=2)
    doc.add_paragraph(
        "OSFI E-23 defines the model lifecycle as five stages, each mapped to specific subcomponents "
        "and Principles."
    )
    lifecycle_rows = []
    for stage in _LIFECYCLE_STAGES:
        component = OSFI_LIFECYCLE_COMPONENTS.get(stage, {})
        lifecycle_rows.append((
            component.get("name", stage.capitalize()),
            ", ".join(component.get("subcomponents", [])) or "N/A",
            ", ".join(component.get("principles", [])) or "N/A",
            component.get("description", "N/A"),
        ))
    _add_two_col_or_three_col_table(doc, ['Lifecycle stage', 'Subcomponents', 'Principles', 'Description'], lifecycle_rows)
    doc.add_paragraph()

    # --- A.4 Model Inventory Information Tracking Expectations ---
    doc.add_heading('A.4 Model Inventory Information Tracking Expectations (Appendix 1)', level=2)
    doc.add_paragraph(
        "OSFI E-23 Appendix 1 identifies required and optional model inventory tracking fields, and "
        "fields expected to be current at each lifecycle stage."
    )
    p = doc.add_paragraph()
    p.add_run('Required fields: ').bold = True
    p.add_run(", ".join(APPENDIX_1_REQUIRED_FIELDS))

    p = doc.add_paragraph()
    p.add_run('Optional fields: ').bold = True
    p.add_run(", ".join(APPENDIX_1_OPTIONAL_FIELDS))
    doc.add_paragraph()

    stage_field_rows = [
        (stage.capitalize(), ", ".join(fields))
        for stage, fields in APPENDIX_1_STAGE_SPECIFIC.items()
    ]
    _add_two_col_or_three_col_table(doc, ['Lifecycle stage', 'Stage-specific fields to keep current'], stage_field_rows)


# =============================================================================
# ANNEX D: CONFIGURABLE GOVERNANCE MATRIX (full 5-stage matrix, moved out of
# Annex A - this is methodology-generated implementation guidance, not
# official OSFI E-23 text)
# =============================================================================

def _add_annex_d_configurable_governance_matrix(doc: Document):
    doc.add_heading('ANNEX D. CONFIGURABLE GOVERNANCE MATRIX', level=1)
    _add_source_label(doc, "institution_configurable")

    p = doc.add_paragraph()
    p.add_run(
        "This matrix is methodology-generated implementation guidance. It is not verbatim OSFI E-23 "
        "text and does not determine approval, readiness, or blocking status. Institutions must "
        "configure it to their own model risk framework, governance authorities, policies, and risk "
        "appetite."
    ).italic = True
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.add_run(
        "Full lifecycle governance matrix for reference, across all 5 lifecycle stages and all 4 risk "
        "levels. Values reflect only configured lifecycle governance data - unavailable values display "
        "'Not configured.' rather than being inferred."
    ).italic = True
    p.paragraph_format.space_after = Pt(10)

    for idx, stage in enumerate(_LIFECYCLE_STAGES, 1):
        doc.add_heading(f'D.{idx} {stage.capitalize()}', level=2)
        comparison = get_lifecycle_requirements_comparison(stage)
        if not comparison:
            doc.add_paragraph("Not configured.")
            continue

        rows = []
        for req_name, by_level in comparison.items():
            req_display = req_name.replace('_', ' ').title()
            rows.append((
                req_display,
                by_level.get('low') or "Not configured.",
                by_level.get('medium') or "Not configured.",
                by_level.get('high') or "Not configured.",
                by_level.get('critical') or "Not configured.",
            ))
        _add_two_col_or_three_col_table(doc, ['Requirement', 'Low', 'Medium', 'High', 'Critical'], rows)
        doc.add_paragraph()


# =============================================================================
# ANNEX E: DETAILED MODEL TYPE CLASSIFICATION EVIDENCE
# =============================================================================
# Full per-check evidence detail for model-type classification, moved out of
# the Section 1.2 main body. These are methodology-specific capability
# evidence checks, not OSFI E-23 lifecycle promotion gates.

def _add_annex_e_model_type_classification_evidence(doc: Document, assessment_results: Dict[str, Any]):
    doc.add_heading('ANNEX E. DETAILED MODEL TYPE CLASSIFICATION EVIDENCE', level=1)
    _add_source_label(doc, "methodology_interpretation", "model_evidence")

    p = doc.add_paragraph()
    p.add_run(
        "Model-type classification evidence checks are methodology-specific checks used to identify the "
        "model's capabilities. They are not lifecycle-stage approval gates and do not determine approval, "
        "readiness, or production use."
    ).italic = True
    p.paragraph_format.space_after = Pt(10)

    classification = assessment_results.get("model_type_classification") or {}
    delivery_model = assessment_results.get("delivery_model") or {}
    gates = classification.get("promotion_gates", {})

    def _gate_row(display_name: str, gate: Dict[str, Any], implication: str) -> tuple:
        return (
            display_name,
            "Verified" if gate.get("verified") else "Not Verified",
            "; ".join(gate.get("evidence", [])) or "None",
            "; ".join(gate.get("missing_evidence", [])) or "None",
            implication,
        )

    genai_gate = gates.get("genai_generation", {})
    traditional_evidence = [
        e for e in genai_gate.get("evidence", [])
        if e.startswith("uses_traditional_ml_or_statistical_model")
    ]
    traditional_row = (
        "Traditional / statistical model evidence",
        "Verified" if traditional_evidence else "Not Verified",
        "; ".join(traditional_evidence) or "None",
        "None" if traditional_evidence else "No explicit confirmation of traditional/statistical processing in the project description.",
        "Confirms the Level 1 baseline (traditional/statistical processing); does not by itself indicate any higher level.",
    )

    delivery_label = delivery_model.get("label", "unknown")
    vendor_delivery_row = (
        "Vendor / platform delivery",
        "Verified" if delivery_label != "unknown" else "Not Verified",
        "; ".join(delivery_model.get("evidence", [])) or "None",
        "; ".join(delivery_model.get("missing_evidence", [])) or "None",
        "Determines the delivery model label; does not affect the model type level.",
    )

    rows = [
        traditional_row,
        _gate_row("GenAI generation", genai_gate, "Required for Level 2 and above."),
        _gate_row(
            "Runtime retrieval for GenAI grounding", gates.get("runtime_retrieval", {}),
            "Required (together with GenAI generation) for Level 3.",
        ),
        _gate_row("Tool / action execution", gates.get("tool_or_action_execution", {}), "Required for Level 4."),
        _gate_row("Autonomous operation", gates.get("autonomous_operation", {}), "Required for Level 5."),
        vendor_delivery_row,
    ]

    _add_two_col_or_three_col_table(
        doc,
        ['Capability evidence check', 'Status', 'Evidence summary', 'Missing evidence', 'Classification implication'],
        rows,
    )


# =============================================================================
# ANNEX B: FIT WITH ENTERPRISE RISK MANAGEMENT
# =============================================================================

def _add_annex_b_erm_fit(doc: Document):
    doc.add_heading('ANNEX B: FIT WITH ENTERPRISE RISK MANAGEMENT', level=1)
    _add_source_label(doc, "methodology_interpretation")

    doc.add_paragraph(
        "AI should be understood as a source of risk within the institution's existing Enterprise Risk "
        "Management framework, not as a separate risk universe. The assessment translates AI/model-"
        "specific failure modes into existing risk concepts such as risk appetite, materiality, "
        "likelihood, impact, controls, residual risk, ownership, monitoring, and escalation."
    )

    # --- B.1 Executive framing ---
    doc.add_heading('B.1 Executive Framing', level=2)
    executive_questions = [
        "What could go wrong with this AI/model use case?",
        "Is the risk within the institution's appetite and tolerance?",
        "What controls are required before deployment or continued use?",
        "Who must approve, monitor, and be accountable for residual risk?",
    ]
    for i, question in enumerate(executive_questions, 1):
        p = doc.add_paragraph(f"{i}. {question}")
        p.paragraph_format.left_indent = Inches(0.25)

    # --- B.2 Mapping to existing risk categories ---
    doc.add_heading('B.2 Mapping to Existing Risk Categories', level=2)
    dimension_links = [
        ("Misuse & Unintended Harm Potential", "Conduct risk, compliance risk, operational risk"),
        ("Output Reliability & Integrity", "Model risk, validation risk, performance risk"),
        ("Fairness & Customer Impact", "Conduct risk, legal risk, compliance risk, reputational risk"),
        ("Operational & Security Risk", "Operational risk, technology risk, cyber risk, resilience risk"),
        ("Model Complexity & Opacity", "Model risk, governance risk, validation risk"),
        ("Governance & Oversight", "Enterprise risk, model governance, accountability"),
        ("Data Provenance & Supply Chain Risk", "Data risk, privacy risk, third-party risk, technology risk"),
        ("Systemic & Concentration Risk", "Strategic risk, third-party concentration risk, resilience risk"),
    ]
    _add_two_col_or_three_col_table(doc, ['AI/model risk dimension', 'Existing risk framework link'], dimension_links)
    doc.add_paragraph()

    # --- B.3 Risk treatment outcomes ---
    doc.add_heading('B.3 Risk Treatment Outcomes', level=2)
    treatments = [
        ("Accept", "Residual risk is within appetite and tolerance and can be monitored"),
        ("Mitigate / Reduce", "Additional controls, testing, monitoring, fallback, human review, or scope limits are required"),
        ("Avoid", "The use case should not proceed, or the model should be removed from use"),
        ("Transfer", "Some exposure is shifted through contracts, insurance, indemnities, vendor obligations, or service-level commitments"),
        ("Escalate", "Risk requires approval by a higher authority, such as senior risk committee, executive management, board committee, or regulator-facing governance"),
    ]
    _add_two_col_or_three_col_table(doc, ['Treatment', 'Meaning'], treatments)
    doc.add_paragraph()

    # --- B.4 Documentation gaps ---
    doc.add_heading('B.4 Documentation Gaps', level=2)
    doc.add_paragraph(
        "When information is missing, the assessment flags a documentation gap rather than assuming the "
        "risk is low. Absence of evidence is relevant because the institution may not be able to "
        "demonstrate effective governance without documented, owned, and tested controls."
    )


# =============================================================================
# ANNEX C: DETAILED QUESTION-BY-QUESTION EVIDENCE
# =============================================================================

def _add_annex_c_detailed_evidence(doc: Document, dimension_assessments: Dict[str, Any],
                                   factor_scores: Dict[str, List[Dict[str, Any]]],
                                   validated_extraction: Dict[str, Any]):
    doc.add_heading('ANNEX C: DETAILED QUESTION-BY-QUESTION EVIDENCE', level=1)
    _add_source_label(doc, "model_evidence", "automated_assessment_output")

    p = doc.add_paragraph()
    p.add_run(
        'Full audit trail: all 47 questions grouped by the 8 dimensions, with scoring criteria, '
        'determined value, evidence status, and resulting action for each.'
    )
    p.paragraph_format.space_after = Pt(12)

    extracted_dims = validated_extraction.get("dimensions", {})

    for dim_id in DIMENSION_ORDER:
        dim_info = get_dimension(dim_id)
        if not dim_info:
            continue

        dim_name = dim_info.get("name", dim_id)
        dim_assessment = dimension_assessments.get(dim_id, {})
        dim_risk_level = dim_assessment.get("risk_level", "Not Assessed")
        not_stated_count = dim_assessment.get("not_stated_count", 0)

        doc.add_heading(dim_name, level=2)

        p = doc.add_paragraph()
        p.add_run('Dimension Risk Level: ').bold = True
        run = p.add_run(dim_risk_level)
        run.bold = True
        run.font.color.rgb = _RISK_COLORS.get(dim_risk_level, RGBColor(0, 0, 0))
        if not_stated_count > 0:
            p.add_run(f' ({not_stated_count} factor(s) not stated, defaulted to Medium)')
        p.paragraph_format.space_after = Pt(8)

        factors = get_dimension_factors(dim_id)
        if not factors:
            doc.add_paragraph('No factors defined for this dimension.')
            continue

        dim_factor_scores = factor_scores.get(dim_id, [])
        dim_extracted = extracted_dims.get(dim_id, {})

        rows = []
        for factor in factors:
            factor_id = factor.get("id", "")
            factor_name = factor.get("name", factor_id)
            factor_type = factor.get("type", "qualitative")

            if factor_type == "quantitative":
                thresholds = factor.get("thresholds", {})
                scoring_criteria = "\n".join(
                    f"{level.title()}: {thresholds.get(level, {}).get('description', 'N/A')}"
                    for level in ["low", "medium", "high", "critical"]
                )
            else:
                levels = factor.get("levels", {})
                scoring_criteria = "\n".join(
                    f"{level.title()}: {levels.get(level, 'N/A')}"
                    for level in ["low", "medium", "high", "critical"]
                )

            factor_score_data = next((fs for fs in dim_factor_scores if fs.get("factor_id") == factor_id), None)

            if factor_score_data:
                extracted_value = factor_score_data.get("value")
                is_not_stated = factor_score_data.get("is_not_stated", False)
                is_review_required = factor_score_data.get("is_portfolio_review_required", False)
                is_not_applicable = factor_score_data.get("is_not_applicable", False)
                risk_level = factor_score_data.get("risk_level", "medium")
                evidence = factor_score_data.get("evidence", "")

                if is_review_required:
                    determined_value = "Portfolio Review Required"
                    evidence_status = "Not verified"
                    missing_evidence = "Institution-wide AI/model inventory data unavailable."
                    resulting_action = "Portfolio review required before scoring."
                elif is_not_applicable:
                    determined_value = f"N/A ({risk_level.title()})"
                    evidence_status = "Not applicable"
                    missing_evidence = ""
                    resulting_action = f"Not applicable - scored as {risk_level.title()}."
                elif is_not_stated or extracted_value is None:
                    determined_value = "NOT_STATED"
                    evidence_status = "Not verified"
                    missing_evidence = "No evidence provided in project description for this factor."
                    resulting_action = "Not stated; defaulted to Medium."
                else:
                    determined_value = f"{extracted_value} ({risk_level.title()})"
                    evidence_status = "Verified"
                    missing_evidence = ""
                    resulting_action = "None required."
            else:
                determined_value = "Not Assessed"
                evidence_status = "Not verified"
                evidence = ""
                missing_evidence = "Factor not present in extraction response."
                resulting_action = "Not stated; defaulted to Medium."

            if not evidence:
                dim_factors = dim_extracted.get("factors", {})
                factor_extracted = dim_factors.get(factor_id, {})
                evidence = factor_extracted.get("evidence", "")

            rows.append((
                factor_id, factor_name, scoring_criteria, determined_value,
                evidence_status, evidence or "", missing_evidence, resulting_action,
            ))

        _add_two_col_or_three_col_table(
            doc,
            ['Question ID', 'Factor/question', 'Scoring criteria', 'Determined value',
             'Evidence status', 'Evidence summary', 'Missing evidence', 'Resulting action'],
            rows,
        )
        doc.add_paragraph()


# =============================================================================
# Backwards compatibility wrapper
# =============================================================================

def generate_design_stage_report(
    project_name: str,
    project_description: str,
    assessment_results: Dict[str, Any],
    doc: Document,
    include_methodology_explanation: bool = True
) -> Document:
    """Backwards compatibility wrapper - defaults to Design stage."""
    return generate_osfi_e23_report(
        project_name=project_name,
        project_description=project_description,
        assessment_results=assessment_results,
        doc=doc,
        current_stage="design",
        include_methodology_explanation=include_methodology_explanation
    )
