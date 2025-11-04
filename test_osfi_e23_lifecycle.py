#!/usr/bin/env python3
"""
OSFI E-23 Lifecycle and Terminology Test Suite

Tests the OSFI E-23 lifecycle-focused report generation functionality:
1. Lifecycle stage detection from project descriptions
2. Design stage report structure and content
3. OSFI Principle X.X terminology compliance
4. OSFI Appendix 1 tracking fields inclusion
5. Stage-specific compliance checklists
"""

import os
import sys
import tempfile
from pathlib import Path
from docx import Document

# Add parent directory to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from osfi_e23_structure import (
    detect_lifecycle_stage,
    get_design_stage_checklist,
    OSFI_PRINCIPLES,
    OSFI_OUTCOMES
)
from osfi_e23_report_generators import generate_design_stage_report
from server import MCPServer


def test_lifecycle_stage_detection():
    """Test stage detection from project descriptions."""
    print("\n" + "="*80)
    print("TEST 1: Lifecycle Stage Detection")
    print("="*80)

    test_cases = [
        {
            "description": "Our bank is developing a new credit risk scoring model for retail lending",
            "expected_stage": "design",
            "label": "Design stage (developing)"
        },
        {
            "description": "We are designing an AI-powered fraud detection system",
            "expected_stage": "design",
            "label": "Design stage (designing)"
        },
        {
            "description": "Our model is currently under independent validation and review",
            "expected_stage": "review",
            "label": "Review stage (under review)"
        },
        {
            "description": "The credit model is being validated by our independent review team",
            "expected_stage": "review",
            "label": "Review stage (validating)"
        },
        {
            "description": "We are deploying the risk model to production this quarter",
            "expected_stage": "deployment",
            "label": "Deployment stage (deploying)"
        },
        {
            "description": "Our model is deployed in production and serving customers",
            "expected_stage": "monitoring",
            "label": "Monitoring stage (deployed)"
        },
        {
            "description": "The legacy pricing model is in production and being monitored",
            "expected_stage": "monitoring",
            "label": "Monitoring stage (monitoring)"
        },
        {
            "description": "We are retiring the old credit model and decommissioning it",
            "expected_stage": "decommission",
            "label": "Decommission stage (retiring)"
        },
        {
            "description": "This is a generic model description without stage indicators",
            "expected_stage": "design",
            "label": "Default to design stage"
        }
    ]

    all_passed = True
    for test_case in test_cases:
        detected = detect_lifecycle_stage(test_case["description"])
        passed = detected == test_case["expected_stage"]
        all_passed = all_passed and passed

        status = "✅ PASS" if passed else "❌ FAIL"
        print(f"\n{status}: {test_case['label']}")
        print(f"  Description: {test_case['description'][:60]}...")
        print(f"  Expected: {test_case['expected_stage']}")
        print(f"  Detected: {detected}")

    if all_passed:
        print("\n✅ Test 1 PASSED: All lifecycle stages detected correctly")
    else:
        print("\n❌ Test 1 FAILED: Some stage detections incorrect")

    return all_passed


def test_design_stage_report_structure():
    """Test Design stage report contains correct sections and content."""
    print("\n" + "="*80)
    print("TEST 2: Design Stage Report Structure")
    print("="*80)

    # Create test assessment results
    assessment_results = {
        "risk_score": 72,
        "risk_level": "High",
        "quantitative_score": 60,
        "qualitative_score": 40,
        "model_complexity": "high",
        "data_quality": "medium",
        "business_impact": "high",
        "governance_maturity": "developing"
    }

    # Generate Design stage report
    doc = Document()
    project_name = "Credit Risk Scoring Model - Design Stage Test"
    project_description = "We are developing a new credit risk scoring model for retail lending applications"

    doc = generate_design_stage_report(
        project_name=project_name,
        project_description=project_description,
        assessment_results=assessment_results,
        doc=doc
    )

    # Save to temporary file for inspection
    with tempfile.NamedTemporaryFile(mode='wb', suffix='.docx', delete=False) as tmp_file:
        doc.save(tmp_file.name)
        tmp_path = tmp_file.name

    # Read back and analyze content
    doc = Document(tmp_path)
    full_text = "\n".join([para.text for para in doc.paragraphs])

    # Clean up
    os.unlink(tmp_path)

    # Test assertions
    tests_passed = []

    # 1. Title includes "Design Stage"
    has_design_stage_title = "Design Stage Compliance Report" in full_text
    tests_passed.append(("Title includes 'Design Stage'", has_design_stage_title))

    # 2. Contains OSFI Principles references
    has_principle_refs = "Principle 3.2" in full_text or "Principle 3.3" in full_text
    tests_passed.append(("Contains Principle 3.2 or 3.3 references", has_principle_refs))

    # 3. NO "Section X.X" references (old terminology)
    has_section_refs = "Section 3.2" in full_text or "Section 3.3" in full_text or "Section 1.1" in full_text
    tests_passed.append(("No 'Section X.X' references (old terminology)", not has_section_refs))

    # 4. Contains "Model Information Profile" (OSFI Appendix 1)
    has_model_profile = "Model Information Profile" in full_text
    tests_passed.append(("Contains Model Information Profile (Appendix 1)", has_model_profile))

    # 5. Contains "Design Stage Compliance Checklist"
    has_checklist = "Design Stage Compliance Checklist" in full_text
    tests_passed.append(("Contains Design Stage Compliance Checklist", has_checklist))

    # 6. Contains "Readiness for Model Review Stage"
    has_readiness = "Readiness for Model Review Stage" in full_text or "Review Stage Readiness" in full_text
    tests_passed.append(("Contains Review Stage readiness assessment", has_readiness))

    # 7. Contains professional validation warning
    has_validation_warning = "professional validation" in full_text.lower() or "independent review" in full_text.lower()
    tests_passed.append(("Contains professional validation warning", has_validation_warning))

    # 8. Does NOT contain detailed Review stage checklist items (should only be readiness assessment)
    has_detailed_review_items = "Independent Model Validation Report" in full_text and "Review stage items" in full_text
    tests_passed.append(("No detailed Review stage implementation items", not has_detailed_review_items))

    # Print results
    all_passed = True
    for test_name, passed in tests_passed:
        status = "✅ PASS" if passed else "❌ FAIL"
        print(f"{status}: {test_name}")
        all_passed = all_passed and passed

    if all_passed:
        print("\n✅ Test 2 PASSED: Design stage report structure correct")
    else:
        print("\n❌ Test 2 FAILED: Some report structure issues detected")

    return all_passed


def test_osfi_terminology():
    """Test that reports use OSFI Principle X.X terminology, not Section X.X."""
    print("\n" + "="*80)
    print("TEST 3: OSFI Terminology Compliance")
    print("="*80)

    # Generate a full report and check terminology
    assessment_results = {
        "risk_score": 65,
        "risk_level": "Medium-High",
        "quantitative_score": 55,
        "qualitative_score": 35
    }

    doc = Document()
    doc = generate_design_stage_report(
        project_name="Terminology Test Model",
        project_description="Developing a test model for terminology validation",
        assessment_results=assessment_results,
        doc=doc
    )

    # Save and read back
    with tempfile.NamedTemporaryFile(mode='wb', suffix='.docx', delete=False) as tmp_file:
        doc.save(tmp_file.name)
        tmp_path = tmp_file.name

    doc = Document(tmp_path)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    os.unlink(tmp_path)

    # Count references
    principle_count = full_text.count("Principle ")
    section_count = full_text.count("Section ")

    # Check for specific valid Principle references
    valid_principles = ["1.1", "1.2", "1.3", "2.1", "2.2", "2.3", "3.1", "3.2", "3.3", "3.4", "3.5", "3.6"]
    principle_refs_found = [p for p in valid_principles if f"Principle {p}" in full_text]

    print(f"\nTerminology Analysis:")
    print(f"  'Principle X.X' references: {principle_count}")
    print(f"  'Section X.X' references: {section_count}")
    print(f"  Valid OSFI Principles found: {', '.join(principle_refs_found) if principle_refs_found else 'None'}")

    # Test assertions
    has_principle_refs = principle_count > 0
    no_section_refs = section_count == 0
    has_valid_principles = len(principle_refs_found) > 0

    print(f"\n{'✅ PASS' if has_principle_refs else '❌ FAIL'}: Report uses 'Principle X.X' references")
    print(f"{'✅ PASS' if no_section_refs else '❌ FAIL'}: Report has NO 'Section X.X' references")
    print(f"{'✅ PASS' if has_valid_principles else '❌ FAIL'}: Uses valid OSFI Principle numbers")

    all_passed = has_principle_refs and no_section_refs and has_valid_principles

    if all_passed:
        print("\n✅ Test 3 PASSED: OSFI terminology correct (Principle X.X, not Section X.X)")
    else:
        print("\n❌ Test 3 FAILED: Terminology issues detected")

    return all_passed


def test_appendix_1_fields():
    """Test that OSFI Appendix 1 model tracking fields are included."""
    print("\n" + "="*80)
    print("TEST 4: OSFI Appendix 1 Model Tracking Fields")
    print("="*80)

    # Generate report
    assessment_results = {
        "risk_score": 58,
        "risk_level": "Medium",
        "model_id": "CRSM-2025-001",
        "model_name": "Retail Credit Risk Scoring Model",
        "model_owner": "Chief Risk Officer",
        "model_developer": "Quantitative Analytics Team",
        "business_line": "Retail Banking"
    }

    doc = Document()
    doc = generate_design_stage_report(
        project_name="Appendix 1 Test Model",
        project_description="Testing OSFI Appendix 1 field inclusion",
        assessment_results=assessment_results,
        doc=doc
    )

    with tempfile.NamedTemporaryFile(mode='wb', suffix='.docx', delete=False) as tmp_file:
        doc.save(tmp_file.name)
        tmp_path = tmp_file.name

    doc = Document(tmp_path)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    os.unlink(tmp_path)

    # Check for required Appendix 1 fields
    required_fields = [
        ("Model Information Profile", "Section header"),
        ("Model ID" or "Model Identifier", "Unique identifier"),
        ("Model Name", "Descriptive name"),
        ("Model Owner", "Accountability"),
        ("Model Developer", "Development team"),
        ("Business Line" or "Business Purpose", "Business context")
    ]

    fields_found = []
    for field_name, description in required_fields:
        if isinstance(field_name, tuple):
            found = any(f in full_text for f in field_name)
            field_label = " or ".join(field_name)
        else:
            found = field_name in full_text
            field_label = field_name

        fields_found.append((field_label, description, found))
        status = "✅ PASS" if found else "❌ FAIL"
        print(f"{status}: {field_label} ({description})")

    all_found = all(found for _, _, found in fields_found)

    if all_found:
        print("\n✅ Test 4 PASSED: All required OSFI Appendix 1 fields present")
    else:
        print("\n❌ Test 4 FAILED: Some Appendix 1 fields missing")

    return all_found


def test_stage_specific_checklist():
    """Test that Design stage checklist is stage-specific."""
    print("\n" + "="*80)
    print("TEST 5: Stage-Specific Checklist Content")
    print("="*80)

    # Get Design stage checklist
    design_checklist = get_design_stage_checklist()

    # Verify structure
    expected_categories = [
        "model_rationale",
        "model_data",
        "model_development",
        "planning_for_future",
        "design_governance",
        "readiness_for_review"
    ]

    tests_passed = []

    # 1. Check all expected categories present
    has_all_categories = all(cat in design_checklist for cat in expected_categories)
    tests_passed.append(("All Design stage categories present", has_all_categories))
    print(f"{'✅ PASS' if has_all_categories else '❌ FAIL'}: All Design stage categories present")

    if has_all_categories:
        print(f"  Categories: {', '.join(expected_categories)}")

    # 2. Check that items reference Design stage principles (3.2, 3.3)
    all_items = []
    for category_items in design_checklist.values():
        all_items.extend(category_items)

    design_principle_refs = sum(1 for item in all_items
                                if "3.2" in item.get("osfi_reference", "") or
                                   "3.3" in item.get("osfi_reference", ""))

    has_design_refs = design_principle_refs > 0
    tests_passed.append(("Items reference Design stage principles (3.2, 3.3)", has_design_refs))
    print(f"{'✅ PASS' if has_design_refs else '❌ FAIL'}: Items reference Principles 3.2 or 3.3")
    print(f"  Found {design_principle_refs} references to Design stage principles")

    # 3. Check that items are marked as "design" category
    design_category_items = sum(1 for item in all_items if item.get("category") == "design")
    has_design_category = design_category_items > 0
    tests_passed.append(("Items marked as 'design' category", has_design_category))
    print(f"{'✅ PASS' if has_design_category else '❌ FAIL'}: Items marked as 'design' category")
    print(f"  Found {design_category_items} items with category='design'")

    # 4. Check that items have required fields
    total_items = len(all_items)
    items_with_all_fields = sum(1 for item in all_items
                                if all(field in item for field in ["item", "requirement", "osfi_reference", "category"]))

    all_items_complete = items_with_all_fields == total_items
    tests_passed.append(("All items have required fields", all_items_complete))
    print(f"{'✅ PASS' if all_items_complete else '❌ FAIL'}: All items have required fields")
    print(f"  {items_with_all_fields}/{total_items} items complete")

    # 5. Check for "planning_for_future" category (Design stage planning for future stages)
    has_planning = "planning_for_future" in design_checklist and len(design_checklist["planning_for_future"]) > 0
    tests_passed.append(("Includes planning for future lifecycle stages", has_planning))
    print(f"{'✅ PASS' if has_planning else '❌ FAIL'}: Includes planning for future stages")

    all_passed = all(passed for _, passed in tests_passed)

    if all_passed:
        print("\n✅ Test 5 PASSED: Design stage checklist is properly stage-specific")
    else:
        print("\n❌ Test 5 FAILED: Some checklist structure issues")

    return all_passed


def test_full_export_integration():
    """Test full export integration through MCP server."""
    print("\n" + "="*80)
    print("TEST 6: Full Export Integration with Lifecycle Detection")
    print("="*80)

    server = MCPServer()

    test_cases = [
        {
            "project_name": "Credit Risk Model - Design Phase",
            "project_description": "Our bank is developing a new credit risk scoring model for retail lending",
            "expected_stage": "design",
            "expected_in_filename": "Design_Stage"
        },
        {
            "project_name": "Fraud Detection Model - Under Review",
            "project_description": "Our fraud detection model is currently under independent validation",
            "expected_stage": "review",
            "expected_in_filename": "Review_Stage"
        },
        {
            "project_name": "Pricing Model - Production",
            "project_description": "The pricing model is deployed in production and being monitored",
            "expected_stage": "monitoring",
            "expected_in_filename": "Monitoring_Stage"
        }
    ]

    all_passed = True

    for test_case in test_cases:
        print(f"\nTesting: {test_case['project_name']}")

        assessment_results = {
            "risk_score": 68,
            "risk_level": "Medium-High",
            "quantitative_score": 50,
            "qualitative_score": 40
        }

        result = server._export_e23_report({
            "project_name": test_case["project_name"],
            "project_description": test_case["project_description"],
            "assessment_results": assessment_results
        })

        # Check for success
        if "error" in result:
            print(f"  ❌ FAIL: Export returned error: {result.get('reason', 'Unknown error')}")
            all_passed = False
            continue

        # Check lifecycle stage detected
        detected_stage = result.get("lifecycle_stage")
        stage_correct = detected_stage == test_case["expected_stage"]
        print(f"  {'✅ PASS' if stage_correct else '❌ FAIL'}: Lifecycle stage detected as '{detected_stage}' (expected '{test_case['expected_stage']}')")

        # Check filename includes stage
        file_path = result.get("file_path", "")
        filename = os.path.basename(file_path)
        filename_correct = test_case["expected_in_filename"] in filename
        print(f"  {'✅ PASS' if filename_correct else '❌ FAIL'}: Filename includes stage: {filename}")

        # Check file was created
        file_exists = os.path.exists(file_path) if file_path else False
        print(f"  {'✅ PASS' if file_exists else '❌ FAIL'}: File created successfully")

        # Clean up
        if file_exists:
            os.unlink(file_path)

        test_passed = stage_correct and filename_correct and file_exists
        all_passed = all_passed and test_passed

    if all_passed:
        print("\n✅ Test 6 PASSED: Full export integration working correctly")
    else:
        print("\n❌ Test 6 FAILED: Some integration issues detected")

    return all_passed


def run_all_tests():
    """Run all OSFI E-23 lifecycle tests."""
    print("\n" + "="*80)
    print("OSFI E-23 LIFECYCLE AND TERMINOLOGY TEST SUITE")
    print("="*80)
    print("\nTesting lifecycle-focused report generation with official OSFI terminology")

    results = []

    # Run all tests
    results.append(("Lifecycle Stage Detection", test_lifecycle_stage_detection()))
    results.append(("Design Stage Report Structure", test_design_stage_report_structure()))
    results.append(("OSFI Terminology Compliance", test_osfi_terminology()))
    results.append(("OSFI Appendix 1 Fields", test_appendix_1_fields()))
    results.append(("Stage-Specific Checklist", test_stage_specific_checklist()))
    results.append(("Full Export Integration", test_full_export_integration()))

    # Summary
    print("\n" + "="*80)
    print("TEST SUMMARY")
    print("="*80)

    total = len(results)
    passed = sum(1 for _, result in results if result)
    failed = total - passed

    for test_name, result in results:
        status = "✅ PASS" if result else "❌ FAIL"
        print(f"{status}: {test_name}")

    print(f"\nTotal tests: {total}")
    print(f"Passed: {passed} ✅")
    print(f"Failed: {failed} ❌")

    if failed == 0:
        print("\n🎉 ALL TESTS PASSED! OSFI E-23 lifecycle functionality is working correctly.")
        print("\nWhat this means:")
        print("  1. Lifecycle stages are correctly detected from project descriptions")
        print("  2. Design stage reports use proper OSFI E-23 structure")
        print("  3. Reports use official 'Principle X.X' terminology (not 'Section X.X')")
        print("  4. OSFI Appendix 1 model tracking fields are included")
        print("  5. Design stage checklists are properly stage-specific")
        print("  6. Full export integration works with lifecycle detection")
        return True
    else:
        print(f"\n⚠️  {failed} TEST(S) FAILED - Review output above for details")
        return False


if __name__ == "__main__":
    success = run_all_tests()
    sys.exit(0 if success else 1)
